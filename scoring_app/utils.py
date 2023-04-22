from .models import Calculation, CalculationResult
from django.core.paginator import Paginator, PageNotAnInteger, EmptyPage
from django.db.models import Q
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.shared import Pt, RGBColor, Mm, Cm
import uuid, base64
from io import BytesIO
from matplotlib import pyplot as plt
import pandas as pd

def get_graph():
    buffer = BytesIO()
    plt.savefig(buffer, format='png')
    buffer.seek(0)
    image_png = buffer.getvalue()
    graph = base64.b64encode(image_png)
    graph = graph.decode('utf-8')
    buffer.close()
    return graph

def get_plot(calculation_id, format):
    # Full data
    df = pd.read_csv('credit_risk_dataset.csv')
    df = df.drop(['loan_status', 'loan_percent_income', 'loan_status'], axis=1)
    loan_intent_values = pd.value_counts(df['loan_intent']).values
    loan_intent_labels=pd.value_counts(df['loan_intent']).index
    loan_intent_human_labels = ['Образование', 'Медицина', 'Инвестиции', 'Личная', 'Погашение долга', 'Жилищный ремонт']
    home_ownership_values = pd.value_counts(df['person_home_ownership']).values
    home_ownership_labels=pd.value_counts(df['person_home_ownership']).index.tolist()
    home_ownership_human_labels = ['Аренда', 'Ипотека', 'Собственность', 'Другое']
    loan_int_rate = df['loan_int_rate'].sort_values()
    loan_amnt = df['loan_amnt'].sort_values()
    loan_default_values = pd.value_counts(df['cb_person_default_on_file']).values
    loan_default_labels = pd.value_counts(df['cb_person_default_on_file']).index
    loan_default_human_labels = ['Нет', 'Да']
    history_length_values = pd.value_counts(df['cb_person_cred_hist_length']).values
    history_length_labels = pd.value_counts(df['cb_person_cred_hist_length']).index
    person_age = df['person_age'].sort_values()[:-80]
    loan_grade_values = pd.value_counts(df['loan_grade']).values
    loan_grade_labels = pd.value_counts(df['loan_grade']).index.tolist()
    # pd.value_counts(df['person_income']).values
    calculation = Calculation.objects.get(id = calculation_id)
    calculation_result = CalculationResult.objects.get(calculation_id = calculation_id)

    explode_home_ownership = [0, 0, 0, 0]
    explode_loan_intent = [0, 0, 0, 0, 0, 0]
    explode_loan_grade = [0, 0, 0, 0, 0, 0, 0]
    counter_home_ownership = 0
    counter_loan_intent = 0
    counter_loan_grade = 0
    for home_ownership in home_ownership_labels:
        if home_ownership == calculation.person_home_ownership:
            explode_home_ownership[counter_home_ownership] = 0.1
        counter_home_ownership += 1

    for loan_intent in loan_intent_labels:
        if loan_intent == calculation.loan_intent:
            explode_loan_intent[counter_loan_intent] = 0.1
        counter_loan_intent += 1

    for loan_grade in loan_grade_labels:
        if loan_grade == calculation_result.score:
            explode_loan_grade[counter_loan_grade] = 0.1
        counter_loan_grade += 1
    
    sum_home_ownership = pd.value_counts(df['person_home_ownership']).values.sum()
    sum_loan_intent = pd.value_counts(df['loan_intent']).values.sum()
    sum_loan_grade = pd.value_counts(df['loan_grade']).values.sum()

    counter_home_ownership = 0
    counter_loan_intent = 0
    counter_loan_grade = 0

    percent_home_ownership = []
    percent_loan_intent = []
    percent_loan_grade = []
    for home_ownership_value in home_ownership_values:
        percent_home_ownership.append(str(round(home_ownership_value/sum_home_ownership, 3)*100)+'%')
    for loan_intent_value in loan_intent_values:
        percent_loan_intent.append(str(round(loan_intent_value/sum_loan_intent, 2)*100)+'%')
    for loan_grade_value in loan_grade_values:
        percent_loan_grade.append(str(round(loan_grade_value/sum_loan_grade, 3)*100)+'%')

    for home_ownership_human_label in home_ownership_human_labels:
        if float(percent_home_ownership[counter_home_ownership][:-1]) > 1:
            home_ownership_human_labels[counter_home_ownership] = home_ownership_human_label + f' ({percent_home_ownership[counter_home_ownership]})'
        else:
            home_ownership_human_labels[counter_home_ownership] = home_ownership_human_label + f' (<1%)'
        counter_home_ownership += 1
    for loan_intent_human_label in loan_intent_human_labels:
        if float(percent_loan_intent[counter_loan_intent][:-1]) > 1:
            loan_intent_human_labels[counter_loan_intent] = loan_intent_human_label + f' ({percent_loan_intent[counter_loan_intent]})'
        else:
            loan_intent_human_labels[counter_loan_intent] = loan_intent_human_label + f' (<1%)'
        counter_loan_intent += 1
    for loan_grade_label in loan_grade_labels:
        if float(percent_loan_grade[counter_loan_grade][:-1]) > 1:
            loan_grade_labels[counter_loan_grade] = loan_grade_label + f' ({percent_loan_grade[counter_loan_grade]})'
        counter_loan_grade += 1

    colors = ["#00ebc1", "#00b8ba", "#0085b4", "#0051ad", "#001ea6"]
    wedge_properties = {"edgecolor":"k",'linewidth': 2}
    plt.switch_backend('AGG')
    
    if format == 'word':
        fig, axs = plt.subplots(2, 3, figsize=(18, 10))
        fig.set_facecolor('white')

        axs[0, 0].set_title('Статус владения жильем')
        axs[0, 0].pie(home_ownership_values, labels=home_ownership_human_labels, startangle=30, counterclock=False, colors=colors, wedgeprops=wedge_properties, labeldistance=1.1, explode=explode_home_ownership, shadow=True)
        
        axs[0, 1].set_title('Цель использования')
        axs[0, 1].pie(loan_intent_values, labels=loan_intent_human_labels, startangle=30, counterclock=False, colors=colors, wedgeprops=wedge_properties, labeldistance=1.1, explode=explode_loan_intent, shadow=True)
        
        axs[0, 2].set_title('Кредитный риск-профиль')
        axs[0, 2].pie(loan_grade_values, labels=loan_grade_labels, startangle=30, counterclock=False, colors=colors, wedgeprops=wedge_properties, labeldistance=1.1,explode=explode_loan_grade, shadow=True)

        axs[1, 0].set_title('Процентная ставка')
        axs[1, 0].hist(loan_int_rate, bins=15, edgecolor='black', color=colors[4])
        axs[1, 0].set_facecolor('#f1f0ef')

        axs[1, 1].set_title('Сумма кредита или займа')
        axs[1, 1].hist(loan_amnt, bins=15, edgecolor='black', color=colors[2])
        axs[1, 1].set_facecolor('#f1f0ef')

        axs[1, 2].set_title('Возраст')
        axs[1, 2].hist(person_age, bins=20, edgecolor='black', color=colors[1])
        axs[1, 2].set_facecolor('#f1f0ef')

        #axs[2, 0].set_title('Кредитный дефолт')
        #axs[2, 0].bar(loan_default_human_labels, loan_default_values, edgecolor='black')
        #axs[2, 0].set_facecolor('#f1f0ef')

        plt.savefig('static/files/plot.png')
        plt.tight_layout()
    else:
        fig, axs = plt.subplots(6, 1, figsize=(5, 18), gridspec_kw={'height_ratios': [5, 5, 5, 5, 5, 5]})
        fig.set_facecolor('#f1f0ef')

        axs[0].set_title('Статус владения жильем')
        axs[0].pie(home_ownership_values, labels=home_ownership_human_labels, startangle=30, counterclock=False, colors=colors, wedgeprops=wedge_properties, pctdistance=0.7, explode=explode_home_ownership, shadow=True)
        
        axs[1].set_title('Цель использования')
        axs[1].pie(loan_intent_values, labels=loan_intent_human_labels, startangle=30, counterclock=False, colors=colors, wedgeprops=wedge_properties, pctdistance=0.7, explode=explode_loan_intent, shadow=True)
        
        axs[2].set_title('Кредитный риск-профиль')
        axs[2].pie(loan_grade_values, labels=loan_grade_labels, startangle=30, counterclock=False, colors=colors, wedgeprops=wedge_properties, labeldistance=1.1,explode=explode_loan_grade, shadow=True)

        axs[3].set_title('Процентная ставка')
        axs[3].hist(loan_int_rate, bins=15, edgecolor='black', color=colors[4])
        axs[3].set_facecolor('#f1f0ef')

        axs[4].set_title('Сумма кредита или займа')
        axs[4].hist(loan_amnt, bins=15, edgecolor='black', color=colors[2])
        axs[4].set_facecolor('#f1f0ef')

        axs[5].set_title('Возраст')
        axs[5].hist(person_age, bins=20, edgecolor='black', color=colors[1])
        axs[5].set_facecolor('#f1f0ef')

        #axs[5].set_title('Кредитный дефолт')
        #axs[5].bar(loan_default_human_labels, loan_default_values, edgecolor='black')
        #axs[5].set_facecolor('#f1f0ef')

        plt.tight_layout()
    graph = get_graph()
    return graph

def change_orientation(document):
    current_section = document.sections[-1]
    new_width, new_height = current_section.page_height, current_section.page_width
    new_section = document.add_section(WD_SECTION.NEW_PAGE)
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width = new_width
    new_section.page_height = new_height

    return new_section

def create_word_results(print_query):
    document = Document()
    styles = document.styles

    paragraph_styles = document.styles['Normal']
    paragraph_font = paragraph_styles.font
    paragraph_font.name = 'Times New Roman'
    paragraph_font.size = Pt(14)

    new_heading_style = styles.add_style('New Heading', WD_STYLE_TYPE.PARAGRAPH)
    new_heading_style.base_style = styles['Heading 1']

    heading_font = new_heading_style.font
    heading_font.name = 'Times New Roman'
    heading_font.size = Pt(16)
    heading_font.color.rgb = RGBColor(0, 0, 0)

    heading = document.add_paragraph(f'Ваш кредитный риск профиль\n', style='New Heading')
    heading.add_run(f'{print_query["score"]}')
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    #heading.paragraph_format.space_after = Pt(20)

    if print_query['recommendations']:
        paragraph = document.add_paragraph('')
        paragraph = document.add_paragraph('Рекомендации:')

    for recommendation in print_query['recommendations'].values():
        paragraph = document.add_paragraph(recommendation, style='List Number')
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    new_section = document.add_section(WD_SECTION.CONTINUOUS)
    current_section = document.sections[-1]
    current_section.top_margin = Cm(3)
    current_section.bottom_margin = Cm(2)
    current_section.left_margin = Cm(0)
    current_section.right_margin = Cm(4)
    change_orientation(document)
    picture = document.add_paragraph()
    run = picture.add_run()
    run.add_picture('static/files/plot.png', width=Mm(280), height=Mm(150))
    picture.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.save('static/files/Результат расчета.docx')

def check_recommendations(request, calculation_id):
    
    recommendations = {}
    calculation = Calculation.objects.get(id = calculation_id)
    
    # P = D x F x K, где P - платежеспособность заемщика, D – ежемесячный доход после уплаты налогов, F – срок кредита или займа в месяцах, а К – поправочный коэффициент (если D меньше $500, то К равен 0,3; если D более $500, но менее $1000, то К равен 0,4; если D больше $1000, но менее $2000, то К равен 0,5; если D больше $2000, то К равен 0,6)
    k = 0
    if calculation.person_income < 500:
        k = 0.3
    elif calculation.person_income >= 500 and calculation.person_income < 1000:
        k = 0.4
    elif calculation.person_income >= 1000 and calculation.person_income < 2000:
        k = 0.5
    elif calculation.person_income >= 2000:
        k = 0.6
    p = round(calculation.person_income, 2) * calculation.loan_term * k

    # S = P/(1 + (C x (t + 1))/(2*12*100))
    s =  p/(1+(calculation.loan_int_rate*(calculation.loan_term+1))/(2*12*100))

    if calculation.loan_amnt > s:
        recommendations['loan_amount'] = f'Указанная сумма кредита или займа превышает оптимальную. Оптимальная сумма кредита или займа для вашего расчета = {round(s, 2)}$.'

    if calculation.person_home_ownership == 'MORTGAGE':
        recommendations['mortgage'] = 'Зачастую финансовые организации отказывают в предоставлении кредита по причине имеющейся кредитной нагрузки в виде ипотеки. Рекомендуется закрывать все предыдущие кредиты или займы перед взятием новых.'

    if calculation.cb_person_default_on_file == 'Y':
        recommendations['default'] = 'Наличие просроченных, даже на небольшой срок, долгов является "красным флагом" для финансовых организаций, особенно, банковских организаций. Рекомендуется не запаздывать и не откладывать запланированные выплаты, не имея на это жизненно-важных обоснований, даже если ранее имел место быть просроченный кредит или заём.'
    
    if calculation.loan_int_rate > 11:
        recommendations['loan_int_rate'] = 'Указанная процентная ставка превышает среднее значение. Рекомендуется ознакомиться с продуктами других финансовых организаций самостоятельно или обратиться за помощью к сайтам-агрегаторам.'
    
    if calculation.cb_person_default_on_file == 'DEBTCONSOLIDATION':
        recommendations['loan_default'] = 'Финансовые организации в большинстве случаев не оформляют кредиты и займы для погашения предыдущих задолженностей. Рекомендуется закрывать все предыдущие кредиты или займы перед взятием новых.'

    return recommendations

def paginate_calculations(request, calculations, results):

    page = request.GET.get('page')
    paginator = Paginator(calculations, results)
    try:
        calculations = paginator.page(page)
    except PageNotAnInteger:
        page = 1
        calculations = paginator.page(page)
    except EmptyPage:
        page = paginator.num_pages
        calculations = paginator.page(page)

    leftIndex = (int(page) - 4)
    if leftIndex < 1:
        leftIndex = 1

    rightIndex = (int(page) + 5)
    if rightIndex > paginator.num_pages:
        rightIndex = paginator.num_pages + 1

    custom_range = range(leftIndex, rightIndex)

    return custom_range, calculations, paginator

def search_calculations(request):
    search_query = ''

    if request.GET.get('search_query'):
        search_query = request.GET.get('search_query')

    calculations_results = CalculationResult.objects.all().filter(
        Q(calculation_id__person_name__icontains=search_query) | 
        Q(score__icontains=search_query), 
        user=request.user).order_by('-date_created')

    return search_query, calculations_results
